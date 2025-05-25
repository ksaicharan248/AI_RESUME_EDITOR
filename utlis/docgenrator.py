from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import ast
from data.jd import job_description
from data.reume import rdata
# ==== Load tailored AI response ====







def doc_generator(tailored_data):
    summary = tailored_data["summary"]
    skills = tailored_data["skills"]
    doc = Document(r"C:\Users\saich\Documents\resume_files\K SAI CHARAN.docx") # Load the resume
    # ==== Update Summary (Table 2, Row 2, Cell 0) with 2.5pt spacing ====
    summary_cell = doc.tables[1].rows[1].cells[0]
    summary_cell.text = ""  # Clear old content
    p = summary_cell.paragraphs[0]
    run = p.add_run(summary)
    run.font.size = Pt(12)
    p.paragraph_format.space_before = Pt(2.5)
    p.paragraph_format.space_after = Pt(2.5)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    # ==== Update Skills (Table 4) ====
    table4 = doc.tables[3]

    # Row 2 - Programming & Scripting
    row1_cell = table4.rows[1].cells[1]
    row1_cell.text = ""  # Clear existing
    p1 = row1_cell.paragraphs[0]
    run0 = p1.add_run(", ".join(skills["languages"]))
    run0.font.size = Pt(12)
    p1.paragraph_format.space_before = Pt(3)
    p1.paragraph_format.space_after = Pt(3)
    p1.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Row 3 - Technologies
    row2_cell = table4.rows[2].cells[1]
    row2_cell.text = ""
    p2 = row2_cell.paragraphs[0]
    run1 = p2.add_run(", ".join(skills["technologies"]))
    run1.font.size = Pt(12)
    p2.paragraph_format.space_before = Pt(3)
    p2.paragraph_format.space_after = Pt(3)
    p2.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Row 4 - Tools
    row3_cell = table4.rows[3].cells[1]
    row3_cell.text = ""
    p3 = row3_cell.paragraphs[0]
    run2 = p3.add_run(", ".join(skills["tools"]))
    run2.font.size = Pt(12)
    p3.paragraph_format.space_before = Pt(3)
    p3.paragraph_format.space_after = Pt(3)
    p3.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # ==== Save the updated resume ====
    doc.save(r"./docs/Tailored_Resume.docx")
    return True


